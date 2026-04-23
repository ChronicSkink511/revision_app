from __future__ import annotations

import csv
import io
import json
from datetime import datetime

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas



def _as_structured_payload(result) -> dict:
    return {
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "topics": [
            {
                "topic": t.topic,
                "evidence": t.evidence,
                "notes": t.notes,
                "mcq": t.mcq,
                "short_answer": t.short_answer,
            }
            for t in result.topics
        ],
    }



def _to_json_bytes(result) -> bytes:
    payload = _as_structured_payload(result)
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")



def _to_csv_bytes(result) -> bytes:
    out = io.StringIO()
    writer = csv.writer(out)
    writer.writerow(["topic", "type", "question", "options", "answer", "explanation_or_guide"])

    for topic in result.topics:
        for q in topic.mcq:
            writer.writerow(
                [
                    topic.topic,
                    "mcq",
                    q.get("question", ""),
                    " | ".join(q.get("options", [])),
                    q.get("answer", ""),
                    q.get("explanation", ""),
                ]
            )
        for q in topic.short_answer:
            writer.writerow(
                [
                    topic.topic,
                    "short_answer",
                    q.get("question", ""),
                    "",
                    "",
                    q.get("answer_guide", ""),
                ]
            )

    return out.getvalue().encode("utf-8")



def _to_docx_bytes(result) -> bytes:
    doc = Document()
    doc.add_heading("Revision Notes and Quizzes", level=1)

    for topic in result.topics:
        doc.add_heading(topic.topic, level=2)
        doc.add_paragraph(f"Evidence: {topic.evidence}")
        doc.add_paragraph("Notes:")
        doc.add_paragraph(topic.notes)

        doc.add_paragraph("MCQ:")
        for idx, q in enumerate(topic.mcq, start=1):
            doc.add_paragraph(f"{idx}. {q.get('question', '')}")
            for opt in q.get("options", []):
                doc.add_paragraph(f"- {opt}")
            doc.add_paragraph(f"Answer: {q.get('answer', '')}")

        doc.add_paragraph("Short Answer:")
        for idx, q in enumerate(topic.short_answer, start=1):
            doc.add_paragraph(f"{idx}. {q.get('question', '')}")
            doc.add_paragraph(f"Guide: {q.get('answer_guide', '')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()



def _to_pdf_bytes(result) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40

    def write_line(line: str) -> None:
        nonlocal y
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(40, y, line[:120])
        y -= 14

    write_line("Revision Notes and Quizzes")
    write_line("=" * 40)

    for topic in result.topics:
        write_line(f"Topic: {topic.topic}")
        write_line(f"Evidence: {topic.evidence}")
        write_line("Notes:")
        for line in topic.notes.splitlines():
            write_line(f"  {line}")

        write_line("MCQ:")
        for idx, q in enumerate(topic.mcq, start=1):
            write_line(f"  {idx}. {q.get('question', '')}")
            for opt in q.get("options", []):
                write_line(f"     - {opt}")
            write_line(f"     Answer: {q.get('answer', '')}")

        write_line("Short Answer:")
        for idx, q in enumerate(topic.short_answer, start=1):
            write_line(f"  {idx}. {q.get('question', '')}")
            write_line(f"     Guide: {q.get('answer_guide', '')}")
        write_line("-" * 40)

    c.save()
    return buffer.getvalue()



def export_bundle(result, export_format: str) -> tuple[bytes, str, str]:
    fmt = export_format.lower()
    if fmt == "json":
        return _to_json_bytes(result), "application/json", "revision_bundle.json"
    if fmt == "csv":
        return _to_csv_bytes(result), "text/csv", "revision_bundle.csv"
    if fmt == "docx":
        return (
            _to_docx_bytes(result),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "revision_bundle.docx",
        )
    if fmt == "pdf":
        return _to_pdf_bytes(result), "application/pdf", "revision_bundle.pdf"

    raise ValueError(f"Unsupported export format: {export_format}")
